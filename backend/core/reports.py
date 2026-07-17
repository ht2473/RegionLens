"""Генерация отчётов по региону: xlsx (openpyxl) и docx (python-docx).

Чистое форматирование: на вход — словарь дашборда региона (`queries.region_dashboard`),
на выход — байты файла. Обращений к БД здесь нет (их делает вьюха экспорта) — это держит
модуль тестируемым и сохраняет разделяемость слоёв.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font

# Русские названия доменов индекса (порядок задаёт queries.INDEX_DOMAINS).
_DOMAIN_RU = {
    "economy": "Экономика",
    "income": "Доходы",
    "demography": "Демография",
    "labor": "Труд",
    "infrastructure": "Инфраструктура",
    "health_edu": "Здоровье и образование",
}

# Атрибуция автора в отчёте (требование Приложения 10; совпадает с подвалом сайта).
_AUTHOR = "Кузьмин Евгений Олегович · студенческий билет № 70232275"


def _num(value: Any, prec: int = 1) -> str:
    """Форматировать число с заданной точностью; None → прочерк."""
    return "—" if value is None else f"{value:.{prec}f}"


def region_xlsx(data: dict[str, Any]) -> bytes:
    """Собрать отчёт региона в xlsx и вернуть его байты."""
    name = data.get("region_name") or data["okato"]
    index = data.get("index") or {}

    wb = Workbook()
    ws = wb.active
    ws.title = "Регион"

    title = ws.cell(row=1, column=1, value=f"Отчёт по региону: {name}")
    title.font = Font(bold=True, size=14)
    ws.append(["ОКАТО", data["okato"]])
    ws.append(["Год", data["year"]])
    ws.append(["Федеральный округ", data.get("federal_district") or "—"])

    ws.append([])
    ws.append(["Индекс развития (0–100)", _num(index.get("total_score"))])
    rank = data.get("rank")
    if rank:
        ws.append(["Ранг по индексу", f"{rank['rank']} из {rank['of']}"])
    cluster = data.get("cluster")
    if cluster:
        ws.append(
            ["Тип (кластер)", f"{cluster.get('cluster_id')} — {cluster.get('cluster_label')}"]
        )

    ws.append([])
    header = ws.cell(row=ws.max_row + 1, column=1, value="Баллы по доменам")
    header.font = Font(bold=True)
    ws.append(["Домен", "Балл", "Предыдущий год", "Дельта"])
    for d in index.get("domains", []):
        ws.append(
            [
                _DOMAIN_RU.get(d["domain"], d["domain"]),
                _num(d.get("score")),
                _num(d.get("score_prev")),
                _num(d.get("delta")),
            ]
        )

    ws.append([])
    header = ws.cell(row=ws.max_row + 1, column=1, value="Ключевые метрики (вклад в тип, SHAP)")
    header.font = Font(bold=True)
    ws.append(["Метрика", "Вклад"])
    for s in data.get("shap_top", []):
        ws.append([s.get("metric_name"), _num(s.get("shap_value"), 3)])

    ws.append([])
    ws.append(["Автор", _AUTHOR])
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 22

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def region_docx(data: dict[str, Any]) -> bytes:
    """Собрать отчёт региона в docx и вернуть его байты."""
    name = data.get("region_name") or data["okato"]
    index = data.get("index") or {}

    doc = Document()
    doc.add_heading(f"Отчёт по региону: {name}", level=0)
    doc.add_paragraph(
        f"ОКАТО: {data['okato']}    Год: {data['year']}    "
        f"Федеральный округ: {data.get('federal_district') or '—'}"
    )

    doc.add_heading("Индекс развития", level=1)
    doc.add_paragraph(f"Итоговый индекс (0–100): {_num(index.get('total_score'))}")
    rank = data.get("rank")
    if rank:
        doc.add_paragraph(f"Ранг по индексу: {rank['rank']} из {rank['of']}")
    cluster = data.get("cluster")
    if cluster:
        doc.add_paragraph(
            f"Тип (кластер): {cluster.get('cluster_id')} — {cluster.get('cluster_label')}"
        )

    domains = index.get("domains", [])
    if domains:
        doc.add_heading("Баллы по доменам", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        head = table.rows[0].cells
        head[0].text, head[1].text, head[2].text, head[3].text = (
            "Домен",
            "Балл",
            "Предыдущий год",
            "Дельта",
        )
        for d in domains:
            cells = table.add_row().cells
            cells[0].text = _DOMAIN_RU.get(d["domain"], d["domain"])
            cells[1].text = _num(d.get("score"))
            cells[2].text = _num(d.get("score_prev"))
            cells[3].text = _num(d.get("delta"))

    shap = data.get("shap_top", [])
    if shap:
        doc.add_heading("Ключевые метрики (вклад в тип, SHAP)", level=1)
        for s in shap:
            doc.add_paragraph(
                f"{s.get('metric_name')}: {_num(s.get('shap_value'), 3)}", style="List Bullet"
            )

    doc.add_paragraph()
    doc.add_paragraph(f"Автор: {_AUTHOR}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_html(data: dict[str, Any]) -> str:
    """Собрать HTML отчёта региона для рендера в PDF. Все динамические значения экранируются."""
    from html import escape

    name = data.get("region_name") or data["okato"]
    index = data.get("index") or {}

    meta = (
        f"ОКАТО: {escape(str(data['okato']))} · Год: {escape(str(data['year']))} · "
        f"Федеральный округ: {escape(str(data.get('federal_district') or '—'))}"
    )

    facts = [f"Итоговый индекс (0–100): <strong>{_num(index.get('total_score'))}</strong>"]
    rank = data.get("rank")
    if rank:
        facts.append(f"Ранг по индексу: {escape(str(rank['rank']))} из {escape(str(rank['of']))}")
    cluster = data.get("cluster")
    if cluster:
        facts.append(
            f"Тип (кластер): {escape(str(cluster.get('cluster_id')))} — "
            f"{escape(str(cluster.get('cluster_label') or '—'))}"
        )
    facts_html = "".join(f"<li>{f}</li>" for f in facts)

    domains = index.get("domains", [])
    domain_rows = "".join(
        "<tr><td>{d}</td><td>{s}</td><td>{p}</td><td>{dl}</td></tr>".format(
            d=escape(_DOMAIN_RU.get(row["domain"], row["domain"])),
            s=_num(row.get("score")),
            p=_num(row.get("score_prev")),
            dl=_num(row.get("delta")),
        )
        for row in domains
    )
    domains_html = (
        "<h2>Баллы по доменам</h2>"
        "<table><thead><tr><th>Домен</th><th>Балл</th><th>Предыдущий год</th>"
        f"<th>Дельта</th></tr></thead><tbody>{domain_rows}</tbody></table>"
        if domains
        else ""
    )

    shap = data.get("shap_top", [])
    shap_items = "".join(
        f"<li>{escape(str(s.get('metric_name') or ''))}: {_num(s.get('shap_value'), 3)}</li>"
        for s in shap
    )
    shap_html = (
        f"<h2>Ключевые метрики (вклад в тип, SHAP)</h2><ul>{shap_items}</ul>" if shap else ""
    )

    return f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8"><style>
  @page {{ size: A4; margin: 20mm 18mm; }}
  body {{
    font-family: "DejaVu Sans", sans-serif; color: #1c2530;
    font-size: 11pt; line-height: 1.45;
  }}
  h1 {{ font-size: 20pt; margin: 0 0 4px; }}
  h2 {{
    font-size: 13pt; margin: 18px 0 6px;
    border-bottom: 1px solid #c9cfd6; padding-bottom: 3px;
  }}
  .meta {{ color: #5b6672; font-size: 10pt; margin-bottom: 10px; }}
  ul {{ margin: 4px 0; padding-left: 18px; }}
  table {{ border-collapse: collapse; width: 100%; margin-top: 4px; }}
  th, td {{ border: 1px solid #c9cfd6; padding: 5px 8px; text-align: left; font-size: 10pt; }}
  th {{ background: #eef2f4; }}
  .footer {{
    margin-top: 22px; color: #5b6672; font-size: 9pt;
    border-top: 1px solid #c9cfd6; padding-top: 6px;
  }}
</style></head><body>
  <h1>Отчёт по региону: {escape(str(name))}</h1>
  <p class="meta">{meta}</p>
  <h2>Индекс развития</h2>
  <ul>{facts_html}</ul>
  {domains_html}
  {shap_html}
  <p class="footer">Автор: {escape(_AUTHOR)}</p>
</body></html>"""


def region_pdf(data: dict[str, Any]) -> bytes:
    """Собрать отчёт региона в pdf (weasyprint) и вернуть его байты.

    weasyprint импортируется лениво: на некоторых платформах (например, Windows без GTK)
    его системные библиотеки не установлены, и это не должно ломать импорт модуля и другие
    форматы экспорта. Вьюха экспорта ловит ошибку импорта и отдаёт понятное сообщение.
    """
    from weasyprint import HTML  # локальный импорт — см. docstring

    return bytes(HTML(string=_pdf_html(data)).write_pdf())
