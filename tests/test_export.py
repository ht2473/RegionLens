"""Тесты отчётов и экспорта: валидность xlsx/docx, эндпойнт скачивания, ExportJob.

Отчёты тестируются прямым вызовом на образце данных (форма как у queries.region_dashboard).
Эндпойнт — с подменой источника данных (monkeypatch), чтобы не зависеть от DuckDB; файлы
экспорта пишутся во временный MEDIA_ROOT.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

import pytest
from core import reports
from core.models import ExportJob
from core.views import _validated_okato
from django.contrib.auth.models import User
from django.http import Http404
from django.test import Client
from docx import Document
from openpyxl import load_workbook

pytestmark = pytest.mark.django_db

_PW = "Sl0transit-9"

_SAMPLE: dict[str, Any] = {
    "okato": "45000000",
    "year": 2024,
    "region_name": "Москва",
    "federal_district": "Центральный",
    "index": {
        "total_score": 87.3,
        "total_score_prev": 85.0,
        "total_delta": 2.3,
        "domains": [
            {"domain": "economy", "score": 1.2, "score_prev": 1.0, "delta": 0.2},
            {"domain": "income", "score": 0.9, "score_prev": 0.8, "delta": 0.1},
        ],
    },
    "cluster": {"cluster_id": 2, "cluster_label": "высокие доходы", "distance_to_centroid": 0.5},
    "shap_top": [{"metric_id": 1, "metric_name": "Среднедушевые доходы", "shap_value": 0.123}],
    "rank": {"rank": 1, "of": 85},
}


def test_region_xlsx_is_valid_workbook() -> None:
    """xlsx — валидная книга (ZIP-сигнатура) с именем региона и доменом внутри."""
    content = reports.region_xlsx(_SAMPLE)
    assert content[:2] == b"PK"
    ws = load_workbook(BytesIO(content)).active
    text = " ".join(str(c.value) for row in ws.iter_rows() for c in row if c.value is not None)
    assert "Москва" in text
    assert "Экономика" in text


def test_region_docx_is_valid_document() -> None:
    """docx — валидный документ с именем региона, доменом и атрибуцией автора."""
    content = reports.region_docx(_SAMPLE)
    assert content[:2] == b"PK"
    doc = Document(BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "Москва" in text
    assert "Экономика" in table_text
    assert "Кузьмин" in text


def test_export_requires_login(client: Client) -> None:
    """Экспорт доступен только вошедшим."""
    resp = client.get("/regions/45000000/export/?format=xlsx&year=2024")
    assert resp.status_code == 302
    assert "/accounts/login/" in resp["Location"]


@pytest.mark.parametrize("fmt", ["xlsx", "docx"])
def test_export_downloads_and_logs(
    client: Client, monkeypatch: pytest.MonkeyPatch, settings: Any, tmp_path: Any, fmt: str
) -> None:
    """Экспорт отдаёт файл-вложение и фиксирует ExportJob для текущего пользователя."""
    settings.MEDIA_ROOT = str(tmp_path)
    monkeypatch.setattr("core.queries.region_dashboard", lambda okato, year: _SAMPLE)
    user = User.objects.create_user("exporter", password=_PW)
    client.force_login(user)

    resp = client.get(f"/regions/45000000/export/?format={fmt}&year=2024")
    assert resp.status_code == 200
    assert "attachment" in resp["Content-Disposition"]
    assert f".{fmt}" in resp["Content-Disposition"]
    assert b"".join(resp.streaming_content)[:2] == b"PK"

    job = ExportJob.objects.get(user=user)
    assert job.fmt == fmt
    assert job.okato == "45000000"
    assert job.status == ExportJob.Status.DONE
    assert job.file


def test_export_rejects_unknown_format(client: Client) -> None:
    """Неизвестный формат — 404."""
    user = User.objects.create_user("u2", password=_PW)
    client.force_login(user)
    assert client.get("/regions/45000000/export/?format=pdf").status_code == 404


def test_export_404_when_no_data(
    client: Client, monkeypatch: pytest.MonkeyPatch, settings: Any, tmp_path: Any
) -> None:
    """Если данных по региону/году нет — 404, задание не создаётся."""
    settings.MEDIA_ROOT = str(tmp_path)
    monkeypatch.setattr("core.queries.region_dashboard", lambda okato, year: None)
    user = User.objects.create_user("u3", password=_PW)
    client.force_login(user)
    assert client.get("/regions/00000000/export/?format=xlsx&year=2024").status_code == 404
    assert ExportJob.objects.filter(user=user).count() == 0


@pytest.mark.parametrize("bad_okato", ["..", "45..00", "abc", "..\\sneaky", "1234567890123"])
def test_export_rejects_invalid_okato(
    client: Client,
    monkeypatch: pytest.MonkeyPatch,
    settings: Any,
    tmp_path: Any,
    bad_okato: str,
) -> None:
    """Нечисловой/traversal-okato отклоняется (404) до записи файла — защита от обхода пути.

    Источник данных подменён (вернул бы образец), поэтому 404 доказывает именно валидацию
    входа, а не отсутствие данных; ExportJob при этом не создаётся.
    """
    settings.MEDIA_ROOT = str(tmp_path)
    monkeypatch.setattr("core.queries.region_dashboard", lambda okato, year: _SAMPLE)
    user = User.objects.create_user("badokato", password=_PW)
    client.force_login(user)

    resp = client.get(f"/regions/{bad_okato}/export/?format=xlsx&year=2024")
    assert resp.status_code == 404
    assert ExportJob.objects.filter(user=user).count() == 0


def test_validated_okato_helper() -> None:
    """Хелпер валидации пропускает корректный ОКАТО и отклоняет traversal/нечисловое."""
    assert _validated_okato("45000000") == "45000000"
    for bad in ("..", "../etc", "..\\x", "abc", ""):
        with pytest.raises(Http404):
            _validated_okato(bad)
